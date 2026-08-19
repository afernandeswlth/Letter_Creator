"""
Editable Word (.docx) versions of the form letters.

Rather than re-drawing each letter, we fill the *branded* Word templates in
engine/word_templates/ (copied from public/letter-templates/) with the same
field values the PDF uses — so the .docx looks like the letter and stays fully
editable. Filling is done in place: table letters (Approval / Pre-Approval /
Conditional) by matching the label cell and setting the value cell; prose
letters (Commencement / Discharge / Custom) by replacing known placeholder text.

build_form_docx(letter_type, brand, values) -> docx bytes.
"""
import io
import os
import re

from docx import Document

from approval_letter import loan_term_years

HERE = os.path.dirname(__file__)
TEMPLATE_DIR = os.path.join(HERE, 'word_templates')


# --- low-level docx helpers ------------------------------------------------
def _set_para_text(paragraph, text):
    """Replace a paragraph's whole text, keeping the first run's formatting.
    Multi-line text becomes line breaks within the paragraph."""
    runs = paragraph.runs
    lines = text.split('\n')
    if not runs:
        paragraph.add_run(lines[0])
    else:
        runs[0].text = lines[0]
        for r in runs[1:]:
            r.text = ''
    first = paragraph.runs[0]
    for ln in lines[1:]:
        first.add_break()
        first.add_text(ln)


def _replace_in_para(paragraph, old, new):
    """Replace `old` with `new` inside a paragraph. Prefers an in-run swap
    (keeps formatting); falls back to collapsing runs when it spans them."""
    for r in paragraph.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    full = ''.join(r.text for r in paragraph.runs)
    if old in full and paragraph.runs:
        paragraph.runs[0].text = full.replace(old, new)
        for r in paragraph.runs[1:]:
            r.text = ''
        return True
    return False


def _set_cell(cell, text):
    """Set a table cell's value, keeping the value paragraph's formatting and
    supporting multiple lines. Removes any nested tables in the cell first."""
    for tbl in cell.tables:
        tbl._element.getparent().remove(tbl._element)
    p = cell.paragraphs[0]
    # drop any extra paragraphs in the cell
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    _set_para_text(p, text)


def _fill_label_tables(doc, mapping):
    """For every table, wherever a cell equals a known label set the next cell
    to the mapped value (skips blanks so template defaults survive)."""
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for i in range(len(cells) - 1):
                label = cells[i].text.strip().rstrip(':').strip()
                if label in mapping and mapping[label]:
                    _set_cell(cells[i + 1], mapping[label])


def _g(values, key, default=''):
    v = values.get(key)
    v = v.strip() if isinstance(v, str) else v
    return v if v else default


# --- letter-type fills -----------------------------------------------------
_TABLE_LABELS = {
    'Borrower(s)': 'borrowers', 'Mortgagor(s)': 'mortgagors', 'Guarantor(s)': 'guarantors',
    'Product Name': 'productName', 'Loan Account Number(s)': 'loanAccountNumber',
    'Application Reference No.': 'applicationNumber', 'Loan Amount': 'loanAmount',
    'Loan Term': 'loanTerm', 'Interest Rate': 'interestRate', 'Revert Rate': 'revertRate',
    'Monthly Repayment': 'monthlyRepayment', 'Rate Type': 'rateType', 'Repayment Type': 'repaymentType',
    'Annual Facility Fee': 'annualFacilityFee', 'Monthly Fees': 'monthlyFees',
    'Offset Account': 'offsetAccount', 'Redraw Facility': 'redrawFacility',
    'Security Property': 'securityProperty', 'Our Panel Solicitor': 'panelSolicitor',
    'Special Conditions': 'specialConditions', 'Conditional Approval items': 'conditionalItems',
}


def _table_value(field, values):
    v = _g(values, field)
    if not v:
        return ''
    if field == 'loanTerm':
        return loan_term_years(v)
    if field == 'repaymentType' and v == 'Interest Only':
        yrs = _g(values, 'ioYears')
        if yrs:
            return f'Interest Only – {loan_term_years(yrs)}'
    return v


_PLACEHOLDER = re.compile(r'^[\s$#%]*$')


def _clear_placeholders(doc):
    """Blank any leftover template filler ('####', '##%', '$####') in table
    cells that were never filled — real defaults (WLTH, Ocean, $395.00) stay."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.replace('\n', '')
                if '#' in t and _PLACEHOLDER.match(t):
                    _set_cell(cell, '')


def _fill_approval_family(doc, values):
    mapping = {label: _table_value(field, values) for label, field in _TABLE_LABELS.items()}
    _fill_label_tables(doc, mapping)
    _clear_placeholders(doc)


def _fill_commencement(doc, values):
    _fill_label_tables(doc, {
        'Customer Name(s)': _g(values, 'customerNames'),
        'Application Number': _g(values, 'applicationNumber'),
        'Disbursement Total': _g(values, 'disbursementTotal'),
        'Construction Address': _g(values, 'constructionAddress'),
    })
    name = _g(values, 'builderName')
    addr = _g(values, 'builderAddress')
    date = _g(values, 'date')
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == 'Name' and name:
            _set_para_text(p, name)
        elif t == 'Address' and addr:
            _set_para_text(p, addr)
        elif t == 'Date' and date:
            _set_para_text(p, date)
        elif t == 'Dear Name' and name:
            _set_para_text(p, f'Dear {name.split()[0] if name else "there"}')


def _fill_discharge(doc, values):
    name = _g(values, 'recipientName')
    addr = _g(values, 'recipientAddress')
    date = _g(values, 'date')
    product = _g(values, 'productName', 'Ultra')
    accts = _g(values, 'accountNumbers')
    dis_date = _g(values, 'dischargeDate')
    sec = _g(values, 'securityAddress')
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == 'NAME' and name:
            _set_para_text(p, name)
        elif t == 'ADDRESS' and addr:
            _set_para_text(p, addr)
        elif t.startswith('Dear NAME') and name:
            _set_para_text(p, f'Dear {name}')
        elif t.startswith('Security Address') and sec:
            _set_para_text(p, f'Security Address:  {sec}')
    # embedded placeholders across the body
    for p in doc.paragraphs:
        if product != 'Ultra':
            _replace_in_para(p, 'Ultra Loan', f'{product} Loan')
        if accts:
            _replace_in_para(p, '4xxxxxxx', accts)
        if date:
            _replace_in_para(p, '22 January 2026', date)
        if dis_date:
            _replace_in_para(p, '27 February 2026', dis_date)


def _fill_custom(doc, values):
    name = _g(values, 'recipientName')
    addr = _g(values, 'recipientAddress')
    date = _g(values, 'date')
    salutation = _g(values, 'salutation') or name
    body_html = _g(values, 'body')
    sign_off = _g(values, 'signOff', 'Sincerely,')
    sender = _g(values, 'senderName', 'Firstname Lastname')
    title = _g(values, 'senderTitle')

    body_lines = _html_to_lines(body_html)
    done_body = False
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if i == 0 and name:
            _set_para_text(p, name)
        elif t.startswith('98 Shirley Street') and addr:
            _set_para_text(p, addr)
        elif re.match(r'^\d{1,2} \w+ \d{4}$', t) and date:
            _set_para_text(p, date)
        elif t == 'Dear':
            _set_para_text(p, f'Dear {salutation},' if salutation else 'Dear,')
        elif t == 'Sincerely,':
            _set_para_text(p, sign_off)
        elif t.startswith('Firstname Lastname'):
            _set_para_text(p, f'{sender} – {title}' if title else sender)
        elif t.startswith('Lorem ipsum') and not done_body and body_lines:
            _set_para_text(p, '\n'.join(body_lines))
            done_body = True
        elif t.startswith(('Sed rutrum', 'In hac habitasse')):
            _set_para_text(p, '')  # clear the extra lorem paragraphs


def _html_to_lines(html):
    """Flatten the rich-text body HTML to plain lines (paragraph → blank line).
    Word formatting of the body is a later refinement; the text lands correctly."""
    if not html:
        return []
    if not re.search(r'<[a-zA-Z/][^>]*>', html):
        return [ln for ln in html.replace('\r\n', '\n').split('\n')]
    s = re.sub(r'(?i)</(p|div|li)>', '\n', html)
    s = re.sub(r'(?i)<br\s*/?>', '\n', s)
    s = re.sub(r'<[^>]+>', '', s)
    from html import unescape
    s = unescape(s)
    return [ln.strip() for ln in s.split('\n')]


# New CAM template: 8 tables with #### value placeholders. Maps
# (table_index, row, col) of each value cell to the form field that fills it.
_CAM_FILL = {
    (0, 0, 1): 'borrowers', (0, 1, 1): 'exposureAccount', (0, 2, 1): 'mortgageManager',
    (1, 1, 1): 'exposureBalance', (1, 1, 3): 'exposureBalance',
    (1, 2, 1): 'exposureInterestType', (1, 2, 3): 'exposureLoanPurpose',
    (1, 3, 1): 'proposedSecurity', (1, 3, 3): 'proposedLvr',
    (2, 1, 1): 'personalInfo', (2, 2, 1): 'employment', (2, 3, 1): 'rentalIncome',
    (2, 4, 1): 'security', (2, 5, 1): 'lmi', (2, 6, 1): 'ndi',
    # table 3 (Refinance History) is filled dynamically — see _fill_refinance.
    (4, 1, 0): 'liabilities',
    (5, 1, 0): 'creditHistory',
    (6, 1, 0): 'livingCost', (6, 1, 1): 'policyExceptions',
    (7, 1, 0): 'finalAssessment', (7, 3, 0): 'recommendation',
    (7, 4, 1): 'recommendedName', (7, 5, 1): 'recommendedDate',
}


def _refinance_notes(raw):
    """Parse the refinanceNotes field (JSON array of note strings) into a list.
    Always returns at least one (possibly empty) entry."""
    import json
    items = []
    if raw:
        try:
            parsed = json.loads(raw)
            items = [str(x).strip() for x in parsed] if isinstance(parsed, list) else [str(parsed).strip()]
        except (ValueError, TypeError):
            items = [str(raw).strip()]
    return items or ['']


def _shade_cell(cell, fill):
    """Set a table cell's background shading (hex fill, no leading #)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def _set_cell_width(cell, twips):
    """Set a cell's width in twips (dxa)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(twips))
    tcW.set(qn('w:type'), 'dxa')


# Refinance grid column widths (twips) — the left "Refinance N" column matches the
# Background information table's label column (~99pt).
_RF_LEFT, _RF_RIGHT = 1979, 8086


def _fill_refinance(table, refis):
    """Fill the Refinance History grid: one row per refinance, left cell
    "Refinance N" (grey), right cell the notes. Adds/removes rows to fit."""
    import copy
    from docx.oxml.ns import qn
    while len(table.rows) < len(refis):
        table._tbl.append(copy.deepcopy(table.rows[-1]._tr))
    while len(table.rows) > len(refis):
        table._tbl.remove(table.rows[-1]._tr)
    # Narrow the "Refinance N" column to match the Background label column.
    gcols = table._tbl.tblGrid.findall(qn('w:gridCol'))
    if len(gcols) >= 2:
        gcols[0].set(qn('w:w'), str(_RF_LEFT))
        gcols[1].set(qn('w:w'), str(_RF_RIGHT))
    for i, note in enumerate(refis):
        cells = table.rows[i].cells
        _set_cell(cells[0], 'Refinance %d' % (i + 1))
        _set_cell(cells[1], note)
        _shade_cell(cells[0], 'ECEFF1')
        _set_cell_width(cells[0], _RF_LEFT)
        _set_cell_width(cells[1], _RF_RIGHT)


def _fill_cam(doc, values):
    """Fill the Credit Approval Memorandum template (8 tables + #### placeholders)."""
    g = lambda k, d='': (values.get(k) or d)  # noqa: E731
    tables = doc.tables

    # Date — the first date-like body paragraph at the top of the document.
    if g('date'):
        for p in doc.paragraphs[:4]:
            s = p.text.strip()
            if s and s[0].isdigit() and '/' in s:
                _set_para_text(p, g('date'))
                break

    # Fill the mapped value cells.
    for (ti, ri, ci), field in _CAM_FILL.items():
        try:
            _set_cell(tables[ti].rows[ri].cells[ci], g(field))
        except IndexError:
            pass

    # Refinance History (table 3) — dynamic rows: "Refinance N" | notes.
    try:
        _fill_refinance(tables[3], _refinance_notes(g('refinanceNotes')))
    except IndexError:
        pass

    # Clear any remaining #### placeholders (unmapped value cells, e.g. the blank
    # Name row and the extra refinance/liabilities grid cells).
    for t in tables:
        for row in t.rows:
            done = set()
            for cell in row.cells:
                if id(cell._tc) in done:
                    continue
                done.add(id(cell._tc))
                if '####' in cell.text:
                    _set_cell(cell, '')

    # Signature image in the sign-off (table 7, "Signature:" row).
    sig = values.get('recommendedSignature')
    if sig and 'base64,' in sig:
        import base64
        from docx.shared import Pt
        raw = base64.b64decode(sig.split('base64,', 1)[1])
        try:
            tables[7].rows[6].cells[1].paragraphs[0].add_run().add_picture(io.BytesIO(raw), height=Pt(40))
        except Exception:  # noqa: BLE001
            pass


_FILLERS = {
    'approval': _fill_approval_family,
    'credit-approval-memorandum': _fill_cam,
    'pre-approval': _fill_approval_family,
    'conditional-approval': _fill_approval_family,
    'commencement': _fill_commencement,
    'discharge': _fill_discharge,
    'custom': _fill_custom,
}


def _find(lines, prefix, default):
    for i, ln in enumerate(lines):
        if ln.startswith(prefix):
            return i
    return default


def build_welcome_docx(d, brand, dd_bsb, dd_account, smsf_number=None):
    """Welcome letter as editable Word. The branded template is a shell with two
    'COPY + PASTE CLIENT DETAILS HERE' spots and a direct-debit table; we split
    the engine's rendered letter text into the header block and the loan-details
    block (the middle static sections already live in the template) and drop them
    in, then fill the BSB/account cells."""
    import wlth_letter as WL
    lines = WL.render_text(d, brand, dd_bsb, dd_account, smsf_number).split('\n')
    i_intro = _find(lines, 'In order to best assist you', len(lines))
    i_loan = _find(lines, 'Your Loan Facility Details:', len(lines))
    i_sign = _find(lines, 'Yours sincerely', len(lines))
    header = '\n'.join(lines[:i_intro]).strip()
    details = '\n'.join(lines[i_loan:i_sign]).strip()

    suffix = 'mma' if brand == 'mma' else 'wlth'
    doc = Document(os.path.join(TEMPLATE_DIR, f'welcome-{suffix}.docx'))

    blocks = [header, details]
    n = 0
    for p in doc.paragraphs:
        if p.text.strip() == 'COPY + PASTE CLIENT DETAILS HERE':
            _set_para_text(p, blocks[n] if n < len(blocks) else '')
            n += 1
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            lbl = cells[0].text.strip().lower()
            if len(cells) > 1 and 'bsb' in lbl:
                _set_cell(cells[1], dd_bsb or '')
            elif len(cells) > 1 and 'account' in lbl:
                _set_cell(cells[1], dd_account or '')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_form_docx(letter_type, brand, values):
    filler = _FILLERS.get(letter_type)
    if filler is None:
        raise ValueError(f'no Word template for letter type "{letter_type}"')
    suffix = 'mma' if brand == 'mma' else 'wlth'
    path = os.path.join(TEMPLATE_DIR, f'{letter_type}-{suffix}.docx')
    doc = Document(path)
    filler(doc, values or {})
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


if __name__ == '__main__':
    import sys
    import json
    sys.stdout.buffer.write(build_form_docx(sys.argv[1], sys.argv[2], json.loads(sys.argv[3])))
